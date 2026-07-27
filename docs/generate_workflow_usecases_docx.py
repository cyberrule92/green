#!/usr/bin/env python3
"""Generate the Workflow Orchestration capabilities + use-case document (.docx).

Run:  python docs/generate_workflow_usecases_docx.py
Out:  docs/Adaptive_Green_AI_Workflow_Orchestration.docx

Self-contained (only python-docx). Mirrors the other docs/generate_*.py scripts.
"""
from __future__ import annotations

import datetime as _dt
import json as _json
import re as _re
import sys as _sys
from pathlib import Path

# Import the real node registry + seeded templates so every step-by-step
# walkthrough is generated from the actual graph (never drifts from the product).
_sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
import workflows as _WF          # noqa: E402
import workflow_templates as _WT  # noqa: E402

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

GREEN = RGBColor(0x0F, 0x9D, 0x58)
DARK = RGBColor(0x1B, 0x1D, 0x22)
GREY = RGBColor(0x60, 0x67, 0x70)

OUT = Path(__file__).resolve().parent / "Adaptive_Green_AI_Workflow_Orchestration.docx"


# ─────────────────────────────────────────────────────────────────────────────
# Content
# ─────────────────────────────────────────────────────────────────────────────
NODE_CATALOG = [
    ("manual", "Trigger", "Run on demand or from the Run button; emits the supplied input."),
    ("webhook", "Trigger", "Fires on POST /api/workflows/webhook/{id}; request body becomes the trigger output."),
    ("schedule", "Trigger", "Fires on a 5-field cron schedule while the workflow is enabled (autonomous)."),
    ("carbon_window", "Trigger", "Fires when grid carbon drops below a gCO2/kWh threshold — defers work to clean windows."),
    ("error_trigger", "Trigger", "Runs this workflow as another workflow's failure handler; emits the error context."),
    ("llm", "AI action", "Chat completion through the CSS greenest-feasible router; reports real gCO2."),
    ("rag_query", "AI action", "Hybrid dense+sparse retrieval + cross-encoder rerank over the indexed knowledge base."),
    ("agent_task", "AI action", "One carbon-budgeted agentic coding task; escalates only on frozen-test evidence."),
    ("guardrail", "AI action", "NemoGuardrails safety rail (jailbreak/PII/harmful); branches safe / blocked."),
    ("image_gen", "AI action", "Carbon-capped diffusion via pluggable NIM endpoint with graceful placeholder fallback."),
    ("http_request", "I/O", "Call any external HTTP(S) API; templated method/url/headers/body. Optional stored credential is injected as an auth header at dispatch and never appears in run state."),
    ("notify", "I/O", "Send a webhook POST or an email (SMTP_*); a graceful logged no-op when SMTP is unconfigured."),
    ("transform", "Logic", "Build an output object from templated fields ({{ $node.id.field }})."),
    ("if", "Logic", "Branch on a comparison (==, !=, contains, >, <, empty…); handles true / false."),
    ("switch", "Logic", "Multi-way branch: first matching case wins, else the default handle. One out-handle per case."),
    ("filter", "Logic", "Gate a path: forward the input when the condition holds, otherwise prune everything downstream."),
    ("wait", "Logic", "Pause the run for a bounded number of seconds (WF_MAX_WAIT_S), then forward the input."),
    ("set_variables", "Logic", "Write templated values into the run-level $vars store, readable by any later node without an edge."),
    ("carbon_gate", "Logic", "Branch on live grid carbon; handles green / dirty."),
    ("merge", "Logic", "Join parallel branches; forwards merged upstream output."),
    ("subworkflow", "Logic", "Run another saved workflow as a node; optional item list fans out (map/batch)."),
    ("approval", "Logic", "Pause the run for a human to approve or reject (interrupt/resume); branches approved / rejected."),
]

# The catalog above is curated prose, but it must not drift from the shipped
# registry -- the whole point of this document is that it describes the product.
_CATALOG_TYPES = {t for t, _c, _d in NODE_CATALOG}
_REGISTRY_TYPES = set(_WF.NODE_TYPES)
if _CATALOG_TYPES != _REGISTRY_TYPES:
    raise SystemExit(
        "NODE_CATALOG is out of sync with workflows.NODE_TYPES.\n"
        f"  missing from this doc: {sorted(_REGISTRY_TYPES - _CATALOG_TYPES)}\n"
        f"  no longer in the product: {sorted(_CATALOG_TYPES - _REGISTRY_TYPES)}"
    )

PARITY = [
    ("Directed graph of steps", "StateGraph nodes/edges", "Chains / Runnables", "Node + edge graph, topologically executed"),
    ("Conditional routing", "Conditional edges", "RouterChain", "if / carbon_gate / guardrail branch handles"),
    ("Multi-way routing", "Conditional edges (n-way)", "RouterChain", "switch node — one out-handle per case, plus default"),
    ("Filtering a path", "Conditional edge to END", "RunnableBranch", "filter node — prunes everything downstream on a miss"),
    ("Run-scoped variables", "Shared state channel", "Memory", "set_variables writes $vars; survives an approval pause"),
    ("Delays", "—", "—", "wait node, bounded by WF_MAX_WAIT_S"),
    ("Parallel execution", "Parallel supersteps", "RunnableParallel", "Independent nodes run concurrently per superstep"),
    ("Fan-out / map", "Send() / map", "Runnable.map / .batch", "subworkflow node with items list"),
    ("Composition / subgraphs", "Subgraphs", "SequentialChain nesting", "subworkflow node (depth-guarded)"),
    ("Retries & timeouts", "Node retry policy", "Runnable.with_retry", "Per-node retries / backoff / timeout_s"),
    ("Error handling", "Try/except nodes", "with_fallbacks", "Per-node on_error = stop | continue"),
    ("Failure workflow", "—", "—", "graph.settings.on_error_workflow_id + error_trigger, depth-bounded"),
    ("Cancellation", "—", "—", "POST /runs/{id}/cancel — cooperative, between supersteps"),
    ("Credential store", "—", "—", "Encrypted at rest, tenant-scoped, injected at dispatch only"),
    ("Notifications", "—", "—", "notify node (webhook / SMTP email)"),
    ("Tool calling", "ToolNode", "Tools / agents", "http_request, agent_task, rag_query, image_gen"),
    ("Retrieval augmentation", "Retriever node", "RetrievalQA", "rag_query node (hybrid + rerank)"),
    ("State / memory passing", "Shared state", "Memory", "Node outputs + {{ $node }} / {{ $input }} refs"),
    ("Human-in-the-loop", "interrupt()", "—", "approval node — pause run, resume on approve/reject"),
    ("Scheduling / triggers", "External", "External", "Built-in cron + webhook + carbon-window triggers"),
    ("Observability", "LangSmith (SaaS)", "Callbacks", "Per-node status/carbon run rows + HMAC audit log"),
    ("Carbon-aware routing", "—", "—", "Unique: every AI node CSS-routed to greenest feasible model"),
]

# Each use case: (title, industry, scenario, trigger, chain[list], carbon, analog)
USE_CASES = [
    # ── Customer Operations ──
    ("Autonomous support-ticket responder", "Customer Operations",
     "Inbound support tickets are answered automatically, grounded in your own knowledge base, "
     "with safety rails on both the question and the drafted reply.",
     "webhook (new ticket)",
     ["guardrail (input) — block jailbreak / abuse", "rag_query — retrieve KB context",
      "llm — draft answer grounded in retrieved context", "guardrail (output) — PII redaction / safety",
      "http_request — post the reply to the helpdesk"],
     "The llm node is CSS-routed to the greenest model that still clears the accuracy floor; "
     "low-priority tickets can be routed behind a carbon_gate so bulk answering waits for a clean grid.",
     "LangChain RetrievalQA + guardrails; LangGraph tool node with conditional safety edges."),

    ("Tiered SLA-aware routing", "Customer Operations",
     "Premium customers get a high-accuracy model; standard traffic gets the greenest model that meets a lower bar.",
     "webhook (request + tenant tier)",
     ["transform — detect tier from payload", "if — tier == premium ?",
      "llm (premium, accuracy_floor 0.9) [true]", "llm (standard) [false]", "http_request — respond"],
     "Accuracy floor is passed straight into CSS, so 'premium' raises the floor rather than pinning a model — "
     "the router still picks the greenest candidate that satisfies it.",
     "LangGraph conditional edges / RouterChain."),

    ("Multilingual answer pipeline", "Customer Operations",
     "Detect the customer's language, answer from the KB in English, then translate the reply back.",
     "webhook (message)",
     ["llm — detect language + translate to English", "rag_query — retrieve KB",
      "llm — compose answer", "llm — translate answer to source language", "http_request — deliver"],
     "Three light llm hops on the greenest rung typically cost less carbon than one call to a large "
     "multilingual model, and each hop is metered independently.",
     "LangChain SequentialChain."),

    ("Sentiment-based escalation & handoff", "Customer Operations",
     "Angry or high-risk messages are escalated to a human queue; everything else is auto-resolved.",
     "webhook (message)",
     ["llm — classify sentiment + intent", "if — sentiment == negative ?",
      "http_request — create escalation ticket [true]", "llm + http_request — auto-reply [false]"],
     "Classification runs on the smallest capable model; only the escalated minority ever touches a larger model.",
     "LangGraph interrupt / human-in-the-loop routing."),

    # ── Software Engineering & DevOps ──
    ("Auto-fix failing CI", "Software Engineering & DevOps",
     "A red build with its failing tests is handed to the coding agent, which fixes the code under a carbon budget "
     "and opens a pull request only if the frozen tests pass.",
     "webhook (CI failure + test files)",
     ["agent_task — carbon-budgeted, caller-supplied frozen tests", "if — status == completed ?",
      "http_request — open PR [true]", "http_request — notify channel [false]"],
     "The agent optimises carbon-per-successful-completion (not per token) and escalates rungs only on verifier "
     "evidence — the tests are the frozen spec, so it cannot reward-hack its way to a green run.",
     "LangGraph agent loop with a verifier tool."),

    ("Incident triage & auto-remediation", "Software Engineering & DevOps",
     "Alerts are classified by severity and, for high-severity issues, a remediation task is generated and paged.",
     "webhook (monitoring alert)",
     ["llm — classify severity", "if — severity == high ?",
      "agent_task — draft remediation [true]", "http_request — page on-call", "transform — audit record"],
     "Triage classification uses the greenest rung; expensive remediation generation is gated to the minority of "
     "high-severity alerts.",
     "LangChain router + tool-calling agent."),

    ("Nightly security & dependency digest", "Software Engineering & DevOps",
     "Every morning, pull new CVEs/advisories, summarise impact against your stack, and post to Slack.",
     "schedule (cron '0 6 * * *')",
     ["http_request — fetch advisory feeds", "rag_query — match against your dependency KB",
      "llm — summarise impact + prioritise", "http_request — post to Slack"],
     "The whole digest is scheduled for an off-peak morning window; put it behind a carbon_gate to slip it to the "
     "greenest hour of the early morning.",
     "LangChain scheduled agent."),

    ("PR review assistant", "Software Engineering & DevOps",
     "Every pull request gets an automated first-pass review comment.",
     "webhook (PR opened, diff)",
     ["guardrail — screen secrets in diff", "llm — review diff for bugs/style", "http_request — post PR comment"],
     "Diff review runs on a code-capable green rung; large diffs can be chunked and fanned out via a subworkflow map.",
     "LangChain code-review chain."),

    ("Automated release notes", "Software Engineering & DevOps",
     "On tag, turn the git log since the last release into human-readable release notes and publish.",
     "webhook (release tag) or schedule",
     ["http_request — fetch commit log", "llm — group + rewrite into release notes",
      "guardrail — output check", "http_request — publish to docs/site"],
     "One llm summarisation on the greenest rung replaces a manual editorial pass.",
     "LangChain summarisation chain."),

    # ── Sustainability & ESG ──
    ("Carbon-aware batch summarisation", "Sustainability & ESG",
     "A backlog of documents is summarised, but only while the grid is clean — the flagship green use case.",
     "carbon_window (≤ 200 gCO2/kWh)",
     ["subworkflow (items = document list) → map", "  · per doc: rag_query → llm summarise",
      "transform — collect summaries", "http_request — store results"],
     "Nothing runs until the carbon-window trigger fires; the fan-out then processes the whole batch during the "
     "clean window. Total run carbon is summed across every mapped sub-run.",
     "LangChain Runnable.map / .batch under a scheduler."),

    ("Nightly CSRD / ESG report", "Sustainability & ESG",
     "Assemble the day's routing/carbon audit metrics into a narrative ESG report and file it.",
     "schedule (cron '0 2 * * *')",
     ["http_request — pull audit + carbon metrics", "llm — draft narrative report",
      "transform — attach figures", "http_request — email / archive"],
     "Runs in the low-demand overnight window; the report itself documents the carbon the platform saved that day.",
     "LangChain report-generation chain."),

    ("Green-window job orchestration", "Sustainability & ESG",
     "Heavy, deferrable jobs (retraining, re-indexing, bulk generation) are launched only when carbon is low.",
     "carbon_window (≤ threshold)",
     ["carbon_gate — confirm still green", "http_request — kick off training/index job",
      "http_request — notify owners"],
     "Turns 'run this whenever the grid is clean' into a first-class, declarative trigger rather than a cron guess.",
     "External scheduler + LangGraph conditional start."),

    ("Emissions anomaly alerting", "Sustainability & ESG",
     "Watch grid intensity and platform usage; alert when carbon per request spikes.",
     "schedule (every 15 min)",
     ["http_request — fetch grid + usage", "if — carbon_per_req > baseline ?",
      "llm — explain likely cause [true]", "http_request — alert"],
     "A cheap periodic check that only spends an llm call when an anomaly is actually detected.",
     "LangChain monitoring agent."),

    # ── Content & Marketing ──
    ("Carbon-capped image campaign", "Content & Marketing",
     "Generate a batch of campaign images, at reduced denoising steps, only during clean-grid windows.",
     "carbon_window (≤ 250 gCO2/kWh)",
     ["subworkflow (items = prompt list) → map", "  · per prompt: image_gen (carbon-capped steps)",
      "http_request — upload assets"],
     "image_gen trims its step budget above the carbon threshold, and the whole batch is deferred to a green window.",
     "LangChain .batch over an image tool."),

    ("SEO content pipeline with human sign-off", "Content & Marketing",
     "Turn a topic into an on-brand, safety-checked article, then pause for an editor to approve before it "
     "publishes to the CMS.",
     "webhook (topic brief)",
     ["rag_query — brand voice + fact KB", "llm — draft article", "guardrail — output safety",
      "approval — editor reviews the draft (run pauses)",
      "http_request — publish to CMS [approved]", "transform — archive as draft [rejected]"],
     "Grounding in a brand KB lets the greenest rung produce on-brand copy without a giant model; the approval "
     "node parks the run (and its carbon receipt) until a human signs off, then resumes down the chosen branch.",
     "LangGraph interrupt/resume + sequential content chain."),

    ("Social listening digest", "Content & Marketing",
     "Summarise brand mentions with sentiment and push a daily digest to a dashboard.",
     "schedule (cron '0 8 * * *')",
     ["http_request — fetch mentions", "llm — summarise + sentiment", "transform — shape dashboard payload",
      "http_request — post to dashboard"],
     "One scheduled summarisation replaces continuous polling; runs in a single green morning window.",
     "LangChain summarisation chain."),

    ("Personalised outreach at scale", "Content & Marketing",
     "Personalise an email per recipient in a segment and hand off to the email platform.",
     "webhook (segment export)",
     ["subworkflow (items = recipients) → map", "  · per recipient: llm personalise → http_request to ESP"],
     "Fan-out runs each personalisation on the greenest rung; the aggregate carbon of the send is reported as one number.",
     "LangChain Runnable.map."),

    # ── Data & Analytics ──
    ("Document extraction & routing", "Data & Analytics",
     "Classify an incoming document, extract structured fields, and route it to the right downstream system.",
     "webhook (uploaded document)",
     ["llm — extract fields + doc type (JSON)", "if — doc type == invoice ?",
      "http_request — post to AP system [true]", "http_request — post to generic store [false]"],
     "Extraction runs once on a green rung; branching avoids any redundant model calls.",
     "LangChain output parser + RouterChain."),

    ("Lead enrichment & scoring", "Data & Analytics",
     "Enrich an inbound form submission with third-party data, score it, and sync to the CRM.",
     "webhook (form submission)",
     ["http_request — enrichment API", "llm — score fit + intent", "transform — CRM field mapping",
      "http_request — upsert to CRM"],
     "Scoring uses the smallest capable model; the http hops carry no model carbon at all.",
     "LangChain tool-augmented chain."),

    ("Data-quality validation gate", "Data & Analytics",
     "Fetch a dataset, validate its schema, and only load it if valid — otherwise alert.",
     "schedule or webhook",
     ["http_request — fetch dataset", "if — schema valid ?",
      "http_request — load to warehouse [true]", "http_request — alert data team [false]"],
     "Pure logic + I/O path: zero model carbon unless you add an llm to explain failures.",
     "LangGraph validation branch."),

    ("Knowledge-base auto-curation", "Data & Analytics",
     "Nightly, summarise newly added documents and index them so retrieval stays fresh.",
     "schedule (cron '0 3 * * *')",
     ["http_request — list new documents", "subworkflow (items = new docs) → map",
      "  · per doc: llm summarise → http_request index", "transform — report counts"],
     "Curation is deferred to the cleanest overnight window and fanned out across the new documents.",
     "LangChain indexing pipeline."),

    ("Competitive-intelligence brief", "Data & Analytics",
     "Aggregate competitor feeds and produce a concise weekly intelligence brief.",
     "schedule (cron '0 7 * * 1')",
     ["http_request — fetch competitor feeds", "rag_query — internal context",
      "llm — synthesise brief", "http_request — deliver"],
     "A single weekly synthesis on the greenest rung replaces continuous analyst polling.",
     "LangChain research chain."),

    # ── IT / Security / Compliance ──
    ("Compliance guardrail gateway", "IT / Security / Compliance",
     "A reusable front door that screens any inbound content and any outbound response, with an audit trail.",
     "webhook (content)",
     ["guardrail (input)", "if — blocked ?", "http_request — reject + log [true]",
      "llm — process [false]", "guardrail (output)", "transform — audit record"],
     "The gateway itself is a workflow other workflows call via the subworkflow node — one place to enforce policy.",
     "LangChain guardrails / constitutional chain."),

    ("Prompt A/B evaluation harness", "IT / Security / Compliance",
     "Run the same input through two prompts in parallel and log the comparison for evaluation.",
     "manual (or webhook)",
     ["llm — prompt A", "llm — prompt B  (runs in parallel with A)",
      "merge — join both arms", "transform — diff + score", "http_request — log to eval store"],
     "The two arms execute in the same superstep concurrently; the harness reports each arm's carbon so you can "
     "compare quality *and* cost.",
     "LangSmith evaluation / RunnableParallel."),

    ("Multi-agent research (agent-of-agents)", "IT / Security / Compliance",
     "A planner decomposes a question into sub-questions, a research sub-workflow answers each in parallel, and a "
     "synthesiser composes the final report.",
     "manual or webhook (research question)",
     ["llm — plan → list of sub-questions (JSON)", "subworkflow (items = sub-questions) → map",
      "  · per sub-question: rag_query → llm answer", "llm — synthesise final report"],
     "Hierarchical fan-out: the planner and synthesiser run on a capable rung, while the many parallel research "
     "sub-runs use the greenest rung — carbon scales with breadth, not model size.",
     "LangGraph hierarchical / multi-agent supervisor."),
]

API_ENDPOINTS = [
    ("GET", "/api/workflows/node-types", "Node palette (all types + their params) for the builder."),
    ("GET", "/api/workflows", "List the tenant's workflows."),
    ("POST", "/api/workflows", "Create a workflow (graph validated at 400 before save)."),
    ("GET", "/api/workflows/{id}", "Fetch one workflow (graph included)."),
    ("PUT", "/api/workflows/{id}", "Update name / description / enabled / graph."),
    ("DELETE", "/api/workflows/{id}", "Delete a workflow and its runs."),
    ("POST", "/api/workflows/{id}/run", "Manually trigger a run; returns a run id (executes in background)."),
    ("POST", "/api/workflows/webhook/{id}", "Public webhook trigger; JSON body becomes the trigger output."),
    ("GET", "/api/workflows/{id}/runs", "Recent runs for a workflow (summaries)."),
    ("GET", "/api/workflows/runs/{run_id}", "One run: status, per-node states, awaiting approvals, total gCO2."),
    ("POST", "/api/workflows/runs/{run_id}/approve", "Approve/reject a paused human-approval node and resume the run."),
    ("GET", "/api/workflows/runs/{run_id}/receipt", "Carbon receipt: per-node and per-model gCO2, plus an approximate saving vs always-full-cloud."),
    ("POST", "/api/workflows/runs/{run_id}/cancel", "Cooperatively cancel an in-flight run; nodes already running finish first."),
    ("GET", "/api/workflows/credentials", "List the tenant's stored HTTP credentials — id, name and type only, never the secret."),
    ("POST", "/api/workflows/credentials", "Store a credential (bearer | basic | header), encrypted at rest."),
    ("DELETE", "/api/workflows/credentials/{id}", "Delete a stored credential."),
    ("GET", "/api/workflows/templates", f"Gallery of the {len(_WT.TEMPLATES)} seeded, runnable workflow templates (summaries)."),
    ("POST", "/api/workflows/templates/{id}/instantiate", "Copy a template into a new editable workflow for the tenant."),
]

EXAMPLE_JSON = """{
  "name": "Green-gated KB answer",
  "graph": {
    "nodes": [
      { "id": "t",   "type": "webhook" },
      { "id": "safe","type": "guardrail", "params": { "text": "{{ $trigger.question }}", "phase": "input" } },
      { "id": "kb",  "type": "rag_query", "params": { "query": "{{ $trigger.question }}", "top_k": 6 } },
      { "id": "ask", "type": "llm",
        "params": { "prompt": "Use this context:\\n{{ $node.kb.context }}\\n\\nQ: {{ $trigger.question }}" },
        "retries": 2, "timeout_s": 60, "on_error": "continue" },
      { "id": "out", "type": "http_request",
        "params": { "method": "POST", "url": "https://example.com/reply",
                    "body": "{{ $node.ask.text }}" } }
    ],
    "edges": [
      { "source": "t",   "target": "safe" },
      { "source": "safe","target": "kb", "sourceHandle": "safe" },
      { "source": "kb",  "target": "ask" },
      { "source": "ask", "target": "out" }
    ]
  }
}"""


# ─────────────────────────────────────────────────────────────────────────────
# Walkthrough generation — derived from the REAL template graphs + node registry
# so the step-by-step can never drift from what the product actually does.
# ─────────────────────────────────────────────────────────────────────────────
_CATEGORY_GROUP = {"trigger": "Triggers", "ai": "AI", "logic": "Logic", "io": "Input / Output"}
_TRIGGER_TYPES = {"manual", "webhook", "schedule", "carbon_window"}
# carbon/analog/scenario prose keyed by template id. New templates carry their own
# prose; the original 25 fall back to the USE_CASES list (paired positionally).
_DETAILS: dict[str, dict[str, str]] = {}
for _i, _t in enumerate(_WT.TEMPLATES):
    if _t.get("scenario") or _t.get("carbon") or _t.get("analog"):
        _DETAILS[_t["id"]] = {"scenario": _t.get("scenario") or _t["description"],
                              "carbon": _t.get("carbon", ""), "analog": _t.get("analog", "")}
    elif _i < len(USE_CASES):
        _uc = USE_CASES[_i]
        _DETAILS[_t["id"]] = {"scenario": _uc[2], "carbon": _uc[5], "analog": _uc[6]}


def _palette_label(ntype: str) -> str:
    nt = _WF.NODE_TYPES.get(ntype)
    return nt.label if nt else ntype


def _group_of(ntype: str) -> str:
    nt = _WF.NODE_TYPES.get(ntype)
    return _CATEGORY_GROUP.get(nt.category, "Logic") if nt else "Logic"


def _param_label(ntype: str, name: str) -> str:
    nt = _WF.NODE_TYPES.get(ntype)
    if nt:
        for p in nt.params:
            if p["name"] == name:
                return p["label"]
    return name


def _topo_nodes(graph: dict) -> list[dict]:
    nodes = {n["id"]: n for n in graph["nodes"]}
    preds: dict[str, list[str]] = {i: [] for i in nodes}
    for e in graph["edges"]:
        preds[e["target"]].append(e["source"])
    order, seen = [], set()

    def visit(i, stack):
        if i in seen or i in stack:
            return
        for p in preds[i]:
            visit(p, stack | {i})
        seen.add(i)
        order.append(i)

    for i in nodes:
        visit(i, set())
    return [nodes[i] for i in order]


def _trigger_fields(graph: dict) -> list[str]:
    fields: set[str] = set()
    for n in graph["nodes"]:
        for v in (n.get("params") or {}).values():
            if isinstance(v, str):
                fields.update(_re.findall(r"\$trigger\.(\w+)", v))
    return sorted(fields)


def _clip(text: str, limit: int = 300) -> str:
    text = str(text).replace("\n", " ↵ ")
    return text if len(text) <= limit else text[:limit] + "…"


def _param_lines(node: dict) -> list[str]:
    """Human-readable 'set this field to that value' lines for one node."""
    ntype = node["type"]
    params = node.get("params") or {}
    lines: list[str] = []
    for name, value in params.items():
        if value in (None, "", "{}", {}, []):
            continue
        if name == "fields" and isinstance(value, dict):
            for k, v in value.items():
                lines.append(f"Add field  {k}  =  {_clip(v, 160)}")
            continue
        lines.append(f"{_param_label(ntype, name)}:  {_clip(value)}")
    # advanced node-level controls, if present on the template
    for adv in ("retries", "timeout_s", "on_error"):
        if node.get(adv) not in (None, "", 0):
            lines.append(f"Advanced · {adv.replace('_', ' ')}:  {node[adv]}")
    return lines


def _node_ref(node: dict) -> str:
    return f"“{node['id']}”"


def _trigger_summary(graph: dict) -> str:
    trig = next((n for n in graph["nodes"] if n["type"] in _TRIGGER_TYPES), None)
    if not trig:
        return "This workflow has no trigger."
    t = trig["type"]
    if t == "manual":
        return "Starts when you press Run."
    if t == "webhook":
        return ("Starts when your system sends an HTTP POST to "
                "/api/workflows/webhook/{workflow-id}. You can also press Run to test it.")
    if t == "schedule":
        cron = (trig.get("params") or {}).get("cron", "")
        return (f"Runs automatically on the schedule ‘{cron}’ (minute hour day month weekday). "
                "Enable the workflow to activate the schedule. You can also press Run to test it now.")
    if t == "carbon_window":
        thr = (trig.get("params") or {}).get("threshold_g", 250)
        return (f"Runs automatically whenever the electricity grid is clean (≤ {thr} gCO₂/kWh). "
                "Enable the workflow to activate it. You can also press Run to test it now.")
    return ""


def _inputs_example(graph: dict) -> str | None:
    fields = _trigger_fields(graph)
    if not fields:
        return None
    return _json.dumps({f: f"<your {f}>" for f in fields})


def _has_type(graph: dict, ntype: str) -> bool:
    return any(n["type"] == ntype for n in graph["nodes"])


# ─────────────────────────────────────────────────────────────────────────────
# Rendering helpers
# ─────────────────────────────────────────────────────────────────────────────
def _shade(cell, hex_fill):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcpr.append(shd)


def _header_row(table, labels):
    hdr = table.rows[0].cells
    for i, label in enumerate(labels):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
        _shade(hdr[i], "0F9D58")


def _table(doc, cols, widths=None):
    t = doc.add_table(rows=1, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def _body(doc, text, size=10.5, color=None, italic=False, space=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space)
    return p


def _num(doc, n, text, bold_lead=None):
    """A numbered step line: '3. <bold lead> rest'."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run(f"{n}.  ")
    r.bold = True
    r.font.size = Pt(10)
    if bold_lead:
        rl = p.add_run(bold_lead)
        rl.bold = True
        rl.font.size = Pt(10)
    rr = p.add_run(text)
    rr.font.size = Pt(10)
    return p


def _bullet(doc, text, indent=0.55, size=9.5, mono=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run("•  ")
    r.font.size = Pt(size)
    rr = p.add_run(text)
    rr.font.size = Pt(size)
    if mono:
        rr.font.name = "Consolas"
    if color:
        rr.font.color.rgb = color
    return p


def _mini_heading(doc, text, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = color
    return p


def _render_walkthrough(doc, idx, tpl):
    """Full layman walkthrough for one use case, generated from its template graph."""
    graph = tpl["graph"]
    detail = _DETAILS.get(tpl["id"], {})
    doc.add_heading(f"{idx}. {tpl['name']}", level=3)
    _body(doc, detail.get("scenario", tpl["description"]), size=10, space=3)

    # Facts table (trigger / carbon / analog).
    ft = _table(doc, 2, widths=[1.3, 5.0])
    ordered = _topo_nodes(graph)
    flow = "  →  ".join(_node_ref(n).strip("“”") for n in ordered)
    for k, v in [("How it starts", _trigger_summary(graph)),
                 ("Node flow", flow),
                 ("Why it's green", detail.get("carbon", "")),
                 ("LangGraph/LangChain analog", detail.get("analog", ""))]:
        cells = ft.add_row().cells
        cells[0].text = ""
        rr = cells[0].paragraphs[0].add_run(k)
        rr.bold = True
        rr.font.size = Pt(8.5)
        rr.font.color.rgb = GREEN
        cells[1].text = ""
        cells[1].paragraphs[0].add_run(v).font.size = Pt(8.5)
    ft._tbl.remove(ft.rows[0]._tr)

    inputs = _inputs_example(graph)

    # ── Option A: one-click from the gallery ──
    _mini_heading(doc, "Option A — Use the ready-made template (fastest, ~30 seconds)", RGBColor(0x7C, 0x5C, 0xFF))
    step = 1
    _num(doc, step, "in the top tabs, then click the ✨ Browse templates button in the left panel.",
         bold_lead="Open the Workflows tab "); step += 1
    _num(doc, step, f"and click Use template. It opens on the canvas as a new (disabled) workflow.",
         bold_lead=f"In the gallery, under “{tpl['industry']}”, find “{tpl['name']}” "); step += 1
    if _has_type(graph, "http_request"):
        _num(doc, step, "and replace every https://example.com/... address with your real endpoint "
                        "(select the node, edit the URL field on the right).",
             bold_lead="Click each HTTP node "); step += 1
    if _has_type(graph, "rag_query"):
        _num(doc, step, "so the RAG step has documents to retrieve from (Carbon tab → RAG, or the /api/rag/index API).",
             bold_lead="Make sure your knowledge base has documents indexed "); step += 1
    trig = next((n for n in graph["nodes"] if n["type"] in _TRIGGER_TYPES), {})
    if trig.get("type") in ("schedule", "carbon_window"):
        _num(doc, step, "toggle the workflow to Enabled so it can fire on its own. (Leave it disabled while testing.)",
             bold_lead="After you've reviewed it, "); step += 1
    if inputs:
        _num(doc, step, f"and in the input box provide: {inputs}", bold_lead="Press Run "); step += 1
    else:
        _num(doc, step, "to test it now (or wait for the automatic trigger).", bold_lead="Press Run "); step += 1
    _num(doc, step, "each node lights up grey→blue→green as it runs, and the total gCO₂ appears in the toolbar.",
         bold_lead="Watch it execute: "); step += 1
    if _has_type(graph, "approval"):
        _num(doc, step, "click Approve or Reject in the amber banner — the run resumes down the branch you chose.",
             bold_lead="When it pauses for approval, "); step += 1

    # ── Option B: build it by hand ──
    _mini_heading(doc, "Option B — Build it yourself, node by node", GREEN)
    _num(doc, 1, "in the left panel to start with an empty canvas.", bold_lead="Click + New ")
    _num(doc, 2, "From the palette on the left, add these nodes (click a name to drop it on the canvas). "
                 "Select each one and, in the right-hand panel, set its Label and fields exactly as listed:",
         bold_lead="Add the nodes.  ")
    for n in ordered:
        _bullet(doc, f"{_group_of(n['type'])} → {_palette_label(n['type'])}   —   set Label to “{n['id']}”",
                indent=0.55, size=9.5, color=DARK)
        for line in _param_lines(n):
            _bullet(doc, line, indent=0.9, size=9, mono=False)
    _num(doc, 3, "Drag from a node's right-edge dot to the next node's left-edge dot to make each connection:",
         bold_lead="Wire them together.  ")
    for e in graph["edges"]:
        handle = e.get("sourceHandle")
        suffix = f"   (use the “{handle}” output)" if handle else ""
        _bullet(doc, f"“{e['source']}”  →  “{e['target']}”{suffix}", indent=0.9, size=9)
    step = 4
    if _has_type(graph, "http_request"):
        _num(doc, step, "in each HTTP node with your real endpoint.",
             bold_lead="Replace the https://example.com/... URLs "); step += 1
    _num(doc, step, "at the top, then click Save.", bold_lead="Give the workflow a name "); step += 1
    if inputs:
        _num(doc, step, f"and enter your input: {inputs}", bold_lead="Click Run "); step += 1
    else:
        _num(doc, step, "(or enable it to let the trigger fire automatically).", bold_lead="Click Run "); step += 1
    _num(doc, step, "green = done, red = failed, amber = waiting for approval, grey = skipped branch. "
                    "The toolbar shows the run's total gram-CO₂.", bold_lead="Read the result: ")
    _body(doc, "", space=6)


# ─────────────────────────────────────────────────────────────────────────────
# Build
# ─────────────────────────────────────────────────────────────────────────────
def build() -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ── Cover ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Adaptive Green AI")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = GREEN
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Carbon-Aware Workflow Orchestration")
    r.font.size = Pt(18)
    r.font.color.rgb = DARK
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run(f"Capabilities & {len(_WT.TEMPLATES)} Industry Use Cases")
    r.font.size = Pt(13)
    r.italic = True
    r.font.color.rgb = GREY
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"Himanshu Tripathi · Hewlett Packard Enterprise · {_dt.date.today():%B %Y}")
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    doc.add_paragraph()
    _body(doc,
          "This document describes the workflow-orchestration capability built into the Adaptive Green AI "
          "platform — a visual, node-based automation builder in the style of n8n, Make, and Activepieces, with "
          "the orchestration depth of LangGraph/LangChain and one property neither offers: every AI step is routed "
          "to the greenest model that can still do the job, and whole workflows can be deferred to low-carbon grid "
          f"windows. It then walks through {len(_WT.TEMPLATES)} production use cases across "
          f"{len({t['industry'] for t in _WT.TEMPLATES})} industries.",
          size=11)

    # ── 1. What it is ──
    doc.add_heading("1. What it is", level=1)
    _body(doc,
          "A workflow is a directed graph of typed nodes: one or more triggers feed a chain of action and logic "
          "nodes, wired together on a visual React Flow canvas. An asynchronous engine walks the graph in "
          "dependency order, passes each node's output downstream via template references, and accumulates a real "
          "gram-CO2 receipt for the whole run. Workflows are authored in the 'Workflows' tab of the product, or "
          "over a REST API, and are persisted per tenant.")
    _body(doc, "Design principles:", size=10.5, space=2)
    for b in [
        "Carbon-first — AI nodes reuse the platform's Composite Sustainability Score (CSS) router, so each step "
        "picks the greenest feasible model rather than a fixed one.",
        "Deferrable — carbon-window triggers hold an entire workflow until the grid is clean.",
        "Real receipts — carbon is measured per node and summed per run, never estimated post-hoc.",
        "Safe by construction — guardrail nodes screen inputs and outputs; every run is audit-logged.",
        "Multi-tenant — workflows, runs, and quotas are scoped per tenant.",
    ]:
        p = doc.add_paragraph(b, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    # ── 2. Capabilities vs LangGraph / LangChain ──
    doc.add_heading("2. Orchestration capabilities vs. LangGraph / LangChain", level=1)
    _body(doc,
          "The engine is a general-purpose orchestrator. The table maps each capability to its LangGraph and "
          "LangChain equivalent and to how this platform implements it.")
    t = _table(doc, 4, widths=[1.7, 1.6, 1.5, 2.2])
    _header_row(t, ["Capability", "LangGraph", "LangChain", "Adaptive Green AI"])
    for cap, lg, lc, us in PARITY:
        cells = t.add_row().cells
        for i, val in enumerate((cap, lg, lc, us)):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(8.5)
            if i == 0:
                run.bold = True
    _body(doc, "", space=2)
    _body(doc,
          "Execution model: nodes whose inputs are all resolved run together in one 'superstep', and independent "
          "nodes in a superstep execute concurrently (asyncio.gather under a parallelism cap) — the same "
          "superstep model LangGraph uses. Branch nodes gate their downstream edges, so pruned branches never run "
          "and join nodes wait for every incoming arm.",
          italic=True, size=10)

    # ── 3. Node catalog ──
    doc.add_heading("3. Node catalog", level=1)
    t = _table(doc, 3, widths=[1.3, 1.0, 3.7])
    _header_row(t, ["Node", "Category", "What it does"])
    for name, cat, desc in NODE_CATALOG:
        cells = t.add_row().cells
        cells[0].text = ""
        run = cells[0].paragraphs[0].add_run(name)
        run.bold = True
        run.font.size = Pt(9)
        cells[1].text = ""
        cells[1].paragraphs[0].add_run(cat).font.size = Pt(9)
        cells[2].text = ""
        cells[2].paragraphs[0].add_run(desc).font.size = Pt(9)

    # ── 4. Reliability & control ──
    doc.add_heading("4. Reliability & control features", level=1)
    for label, desc in [
        ("Parallel supersteps", "Independent branches execute concurrently, bounded by WF_MAX_PARALLEL."),
        ("Retries & backoff", "Any node may set retries and retry_backoff_s; failures are retried with linear backoff."),
        ("Timeouts", "Per-node timeout_s wraps the handler in asyncio.wait_for."),
        ("Error policy", "on_error = stop (fail the run) or continue (soft-fail; downstream still runs)."),
        ("Sub-workflows & map", "The subworkflow node runs another saved workflow; an items list fans it out (map/batch), depth-guarded."),
        ("Human-in-the-loop", "The approval node pauses the run and snapshots its state (JSON); a reviewer approves or rejects via API/UI and the run resumes down the chosen branch — LangGraph interrupt/resume parity."),
        ("Autonomous scheduler", "A built-in loop fires schedule (cron) and carbon_window triggers — no external cron service required."),
        ("Guardrails & audit", "Guardrail nodes screen content; every run persists per-node status + carbon and an HMAC-signed audit entry."),
        ("Safety limits", "Node count, step budget, nesting depth, HTTP size/timeout, and queue size are all bounded."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{label} — ")
        r.bold = True
        r.font.size = Pt(10)
        r2 = p.add_run(desc)
        r2.font.size = Pt(10)

    # ── 5. Using the builder (primer) ──
    doc.add_heading("5. Using the builder — a plain-English primer", level=1)
    _body(doc,
          "You don't need to write code. A workflow is just boxes (called nodes) joined by arrows. This section "
          "teaches the whole screen once; every use case in Section 6 then reuses it.")

    _mini_heading(doc, "The screen has four areas", DARK)
    for label, desc in [
        ("Left panel", "your saved workflows, a “✨ Browse templates” button, and the node palette "
                        "(the menu of steps you can add, grouped Triggers / AI / Logic / Input-Output)."),
        ("Canvas (middle)", "the big area where your nodes live. Drag nodes to move them; scroll to zoom."),
        ("Right panel", "settings for whichever node you clicked — this is where you type prompts, URLs, etc."),
        ("Top toolbar", "the workflow name, Save, ▶ Run, and the live grid-carbon reading."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{label} — "); r.bold = True; r.font.size = Pt(10)
        r2 = p.add_run(desc); r2.font.size = Pt(10)

    _mini_heading(doc, "The eight things you'll ever do", GREEN)
    for i, (lead, rest) in enumerate([
        ("Open it: ", "click the Workflows tab at the top of the app."),
        ("Add a node: ", "click any name in the left palette — it drops onto the canvas. Every workflow needs "
                         "one Trigger node (how it starts) plus one or more action nodes."),
        ("Configure a node: ", "click it once to select it, then fill in the fields that appear in the right panel. "
                               "Always set its Label to a short name — that's how you'll refer to it."),
        ("Connect two nodes: ", "each node has a small dot on its right edge (output) and left edge (input). "
                                "Drag from one node's right dot to the next node's left dot to draw the arrow."),
        ("Branch nodes: ", "IF, Carbon gate, Guardrail and Approval nodes have two outputs (e.g. true/false, "
                           "green/dirty, safe/blocked, approved/rejected). Drag from the specific labelled dot."),
        ("Reuse earlier data: ", "type {{ $node.LABEL.field }} inside any field to pull a previous node's output — "
                                 "e.g. {{ $node.ask.text }} inserts the answer from the node you labelled ‘ask’. "
                                 "Use {{ $trigger.x }} for the input you provide when running."),
        ("Save & Run: ", "name the workflow, click Save, then ▶ Run. If it needs input, type it as JSON "
                         "(e.g. {\"question\": \"…\"}). Nodes light up as they execute."),
        ("Read the result: ", "each node turns green (done), red (failed), amber (waiting for your approval) or "
                              "grey (a branch that was skipped). The toolbar shows the total grams of CO₂ the run spent."),
    ], 1):
        _num(doc, i, rest, bold_lead=lead)

    _mini_heading(doc, "Node status colours", DARK)
    lg = _table(doc, 5, widths=[1.1, 1.1, 1.1, 1.2, 1.1])
    _header_row(lg, ["Grey", "Blue", "Green", "Amber", "Red"])
    cells = lg.add_row().cells
    for i, v in enumerate(["waiting / skipped", "running now", "finished OK",
                           "needs your approval", "failed"]):
        cells[i].text = ""
        cells[i].paragraphs[0].add_run(v).font.size = Pt(8.5)

    _body(doc, "", space=2)
    _body(doc,
          "Two things to remember: (1) Templates ship with https://example.com placeholder web addresses — swap "
          "them for your real ones. (2) Any step that reads your knowledge base (RAG retrieve) needs documents "
          "indexed first (Carbon tab → RAG). Scheduled and carbon-window workflows only fire once you switch the "
          "workflow to Enabled.", italic=True, size=10)

    # ── 6. Use cases with full walkthroughs ──
    doc.add_heading("6. Seventy-five use cases — with step-by-step build instructions", level=1)
    _body(doc,
          "Each use case below can be built two ways: Option A instantiates the ready-made template in one click; "
          "Option B walks you through building it by hand, node by node, with the exact value for every field. "
          "Both produce the same workflow. Everything here is generated from the shipped templates, so it matches "
          "the product exactly.")

    # Group by industry so each industry heading appears once (templates were
    # appended over time, not sorted).
    groups: dict[str, list] = {}
    for tpl in _WT.TEMPLATES:
        groups.setdefault(tpl["industry"], []).append(tpl)
    idx = 1
    for industry, items in groups.items():
        doc.add_heading(industry, level=2)
        for tpl in items:
            _render_walkthrough(doc, idx, tpl)
            idx += 1

    # ── 7. API ──
    doc.add_heading("7. REST API", level=1)
    t = _table(doc, 3, widths=[0.7, 2.6, 3.0])
    _header_row(t, ["Method", "Path", "Purpose"])
    for method, path, purpose in API_ENDPOINTS:
        cells = t.add_row().cells
        cells[0].text = ""
        cells[0].paragraphs[0].add_run(method).font.size = Pt(8.5)
        cells[1].text = ""
        rr = cells[1].paragraphs[0].add_run(path)
        rr.font.size = Pt(8.5)
        rr.font.name = "Consolas"
        cells[2].text = ""
        cells[2].paragraphs[0].add_run(purpose).font.size = Pt(8.5)

    # ── 8. Example ──
    doc.add_heading("8. Example workflow (JSON)", level=1)
    _body(doc,
          "A webhook-triggered, guardrailed, retrieval-augmented answerer whose LLM node retries twice, times out "
          "at 60s, and soft-fails so a downstream notifier can still fire:")
    code = doc.add_paragraph()
    run = code.add_run(EXAMPLE_JSON)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    _shade_paragraph(code)

    # ── 9. Roadmap ──
    doc.add_heading("9. Roadmap", level=1)
    for b in [
        "Native loop / while node for iterative refinement.",
        "Per-tenant workflow carbon budgets and quota enforcement.",
        "Versioned workflows with run-against-version and rollback.",
        "Streaming run progress over SSE/WebSocket to the canvas.",
    ]:
        p = doc.add_paragraph(b, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.save(str(OUT))
    print(f"wrote {OUT}  ({OUT.stat().st_size // 1024} KB)")


def _shade_paragraph(paragraph):
    """Light grey background behind a code block."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F3F5")
    ppr.append(shd)


if __name__ == "__main__":
    build()
